import re
import os
import base64
from PIL import Image
import pytesseract
from deep_translator import GoogleTranslator
from fpdf import FPDF
from gtts import gTTS
import PyPDF2
import docx
import pandas as pd
from io import StringIO
from functools import lru_cache
import streamlit as st
from ratelimit import limits, sleep_and_retry  # Add this import
from api.groq_api import stream_groq_response, GroqAPIError
from api.openai_api import stream_openai_response
from api.anthropic_api import stream_anthropic_response
from datetime import datetime

# At the top of utils.py with other imports
import logging
import os
import re
from datetime import datetime
from typing import Optional, List

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)

# Import API keys from environment variables
GROQ_API_KEY = os.getenv('GROQ_API_KEY')
OPENAI_API_KEY = os.getenv('GROQ_API_KEY')
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')

# Utility Functions
def validate_prompt(prompt: str):
    if not prompt.strip():
        raise ValueError("Prompt cannot be empty")

def sanitize_input(input_text):
    return re.sub(r"[^\w\s\-.,?!]", "", input_text).strip()

def perform_ocr(file_obj):
    try:
        # Convert file object to PIL Image
        image = Image.open(file_obj)
        if not image:
            return "Error: Unable to process image content"
        # Use pytesseract to extract text from image
        text = pytesseract.image_to_string(image)
        if text.strip():
            return text
        return "No text found in the image"
    except Exception as e:
        return f"Error processing image: {str(e)}"

def process_excel_file(file_obj):
    wb = openpyxl.load_workbook(file)
    data = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            data.append(" | ".join(map(str, row)))
    return "\n".join(data)


def process_ppt_file(file_obj):
    try:
        prs = Presentation(file_obj)
        text = []
        for slide in prs.slides:
            for shape in slide.shapes:
                if hasattr(shape, "text"):
                    text.append(shape.text)
        logger.info(f"Successfully processed PowerPoint file")
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error processing PowerPoint file: {str(e)}", exc_info=True)
        return f"Error processing PowerPoint file: {str(e)}"

def get_max_token_limit(model: str) -> int:
    token_limits = {
        "gpt-4": 8192,
        "gpt-3.5-turbo": 4096,
        "claude-3-opus-latest": 200000,
        "claude-3-sonnet-latest": 200000,
        "claude-3-haiku-20240307": 200000,
        "llama-3.1-70b-versatile": 8192,
        "mixtral-8x7b-32768": 32768
    }
    return token_limits.get(model, 4096)  # Default to 4096 if model not found

def get_model_options(provider):
    model_options = {
        "Groq": [
            "llama-3.1-70b-versatile",
            "llama-3.1-405b-reasoning",
            "llama-3.1-8b-instant",
            "llama3-groq-70b-8192-tool-use-preview",
            "llama3-70b-8192",
            "mixtral-8x7b-32768",
            "gemma2-9b-it",
            "whisper-large-v3",
        ],
        "OpenAI": [
            "gpt-4o-mini",
            "gpt-4o",
            "gpt-3.5-turbo",
        ],
        "Anthropic": [
            "claude-3-haiku-20240307",
            "claude-3-5-sonnet-latest",
            "claude-3-opus-latest",
        ]
    }
    return model_options.get(provider, [])

@lru_cache(maxsize=None)
def get_api_client(provider):
    if provider == "Groq":
        return GROQ_API_KEY
    elif provider == "OpenAI":
        return OPENAI_API_KEY
    elif provider == "Anthropic":
        import anthropic
        return anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    return None

async def stream_llm_response(client, params, messages, provider):
    try:
        if provider == "Groq":
            async for chunk in stream_groq_response(client, params, messages):
                yield chunk
        elif provider == "OpenAI":
            async for chunk in stream_openai_response(client, params, messages):
                yield chunk
        elif provider == "Anthropic":
            async for chunk in stream_anthropic_response(client, params, messages):
                yield chunk
        else:
            yield "Error: Unsupported provider"
    except GroqAPIError as e:
        yield f"Groq API Error: {e}"
    except Exception as e:
        yield f"API Error: {str(e)}"

async def process_chat_input(prompt):
    messages = st.session_state.messages + [{"role": "user", "content": prompt}]
    full_response = ""

    with st.chat_message("assistant"):
        message_placeholder = st.empty()
        client = get_api_client(st.session_state.provider)
        if client is None:
            message_placeholder.error("API Key not set for selected provider!")
            return

        async for chunk in stream_llm_response(client, st.session_state.model_params, messages, st.session_state.provider):
            full_response += chunk
            message_placeholder.markdown(full_response + "▌")
        message_placeholder.markdown(full_response)

        st.session_state.messages.append({"role": "assistant", "content": full_response})

def handle_file_upload(uploaded_file):
    # Handles the uploaded file and returns its text content
    file_handlers = {
        "application/pdf": lambda f: " ".join(
            page.extract_text() for page in PyPDF2.PdfReader(f).pages
        ),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": lambda f: " ".join(
            paragraph.text for paragraph in docx.Document(f).paragraphs
        ),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.document": lambda f: process_excel_file(
            f
        ),
        "application/vnd.ms-powerpoint": lambda f: process_ppt_file(f),
        "text/plain": lambda f: f.getvalue().decode("utf-8"),
        "text/markdown": lambda f: f.getvalue().decode("utf-8"),
        "image/jpeg": lambda f: perform_ocr(f),
        "image/png": lambda f: perform_ocr(f),
    }

    for file_type, handler in file_handlers.items():
        if uploaded_file.type.startswith(file_type):
            return handler(uploaded_file)

    raise ValueError("Unsupported file type")

def text_to_speech(text: str, lang: str):
    try:
        lang_map = {"English": "en", "Tamil": "ta", "Hindi": "hi"}
        lang_code = lang_map.get(lang, "en")
        tts = gTTS(text=text, lang=lang_code)
        audio_file = os.path.join(os.getcwd(), "temp_audio.mp3")
        tts.save(audio_file)
        with open(audio_file, "rb") as f:
            audio_bytes = f.read()
        os.remove(audio_file)
        st.session_state.audio_base64 = base64.b64encode(audio_bytes).decode()
    except Exception as e:
        raise Exception(f"Error in text-to-speech conversion: {str(e)}")

def generate_openai_tts(text: str, voice: str):
    try:
        response = client.audio.create(model="tts-1", input=text, voice=voice)
        audio_bytes = response.content
        st.session_state.audio_base64 = base64.b64encode(audio_bytes).decode()
    except openai.APIError as e:
        logger.error(f"OpenAI API error: {e}")
        raise
    except Exception as e:
        logger.error(f"Error generating TTS with OpenAI: {e}")
        raise


def translate_text(text: str, target_lang: str) -> str:
    if target_lang == "English":
        return text
    translator = GoogleTranslator(source="auto", target=target_lang)
    return translator.translate(text)

# Feedback functions
def save_feedback(feedback: str):
    if "feedback.json" not in os.listdir():
        with open("feedback.json", "w") as f:
            json.dump([], f)

    with open("feedback.json", "r+") as f:
        feedback_data = json.load(f)
        feedback_data.append(feedback)
        f.seek(0)
        json.dump(feedback_data, f)

def handle_file_upload() -> None:
    uploaded_file = st.file_uploader(
        "Upload a file",
        type=["pdf", "docx", "txt", "md", "jpg", "jpeg", "png", "xlsx", "pptx"],
    )
    if uploaded_file:
        try:
            st.session_state.file_content = process_uploaded_file(uploaded_file)
            st.success("File processed successfully")
        except Exception as e:
            st.error(f"Error processing file: {e}")


def reset_current_chat():
    st.session_state.messages = []
    st.session_state.is_file_response_handled = False


# Function to check for internet connectivity
def is_connected():
    try:
        response = requests.get("http://www.google.com", timeout=5)
        return response.status_code == 200
    except requests.ConnectionError:
        return False


# Save chat history to local storage
def save_chat_history_locally():
    chat_data = [
        {"role": msg["role"], "content": msg["content"]}
        for msg in st.session_state.messages
    ]
    with open("local_chat_history.json", "w") as f:
        json.dump(chat_data, f)


# Load chat history from local storage
def load_chat_history_locally():
    if os.path.exists("local_chat_history.json"):
        with open("local_chat_history.json", "r") as f:
            return json.load(f)
    return []
def get_max_token_limit(model):
    if "mixtral-8x7b-32768" in model:
        return 32768
    elif "llama-3.1-70b-versatile-131072" in model:
        return 131072
    elif "gemma2-9b-it" in model:
        return 8192
    return 4096


def get_model_options(provider):
    if provider == "Groq":
        return [
            "llama-3.1-70b-versatile",
            "llama-3.1-405b-reasoning",
            "llama-3.1-8b-instant",
            "llama3-groq-70b-8192-tool-use-preview",
            "llama3-70b-8192",
            "mixtral-8x7b-32768",
            "gemma2-9b-it",
            "whisper-large-v3",
        ]
    elif provider == "OpenAI":
        return [
            "gpt-4o-mini",
            "gpt-4o",
            "gpt-3.5-turbo",
        ]
    return []


# Add cache decorator for API responses
def cache_response(func):
    async def wrapper(*args, **kwargs):
        cache_key = hashkey(args, kwargs)
        if cache_key in response_cache:
            return response_cache[cache_key]

        result = await func(*args, **kwargs)
        response_cache[cache_key] = result
        return result

    return wrapper


# Rate limiting
@sleep_and_retry
@limits(calls=5, period=60)
async def rate_limited_api_call(client, params, messages):
    return await async_stream_llm_response(client, params, messages)


# Utility Functions
def validate_prompt(prompt: str):
    if not prompt.strip():
        raise ValueError("Prompt cannot be empty")


def sanitize_input(input_text):
    return re.sub(r"[^\w\s\-.,?!]", "", input_text).strip()


def process_uploaded_file(uploaded_file):
    # Handles the uploaded file and returns its text content
    file_handlers = {
        "application/pdf": lambda f: " ".join(
            page.extract_text() for page in PyPDF2.PdfReader(f).pages
        ),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": lambda f: " ".join(
            paragraph.text for paragraph in docx.Document(f).paragraphs
        ),
        "application/vnd.openxmlformats-officedocument.spreadsheetml.document": lambda f: process_excel_file(
            f
        ),
        "application/vnd.ms-powerpoint": lambda f: process_ppt_file(f),
        "text/plain": lambda f: f.getvalue().decode("utf-8"),
        "text/markdown": lambda f: f.getvalue().decode("utf-8"),
        "image/jpeg": lambda f: perform_ocr(f),
        "image/png": lambda f: perform_ocr(f),
    }

    for file_type, handler in file_handlers.items():
        if uploaded_file.type.startswith(file_type):
            return handler(uploaded_file)

    raise ValueError("Unsupported file type")


def perform_ocr(image_file):
    """Perform OCR on the uploaded image file."""
    try:
        image = Image.open(image_file)
        text = pytesseract.image_to_string(image)
        return text
    except Exception as e:
        logger.error(f"Error performing OCR on the image: {e}")
        raise ValueError(
            "OCR processing failed. Please ensure the image is valid and readable."
        )


def process_excel_file(file):
    wb = openpyxl.load_workbook(file)
    data = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        for row in ws.iter_rows(values_only=True):
            data.append(" | ".join(map(str, row)))
    return "\n".join(data)


def process_ppt_file(file):
    try:
        prs = Presentation(file)
        text = []
        for slide in prs.slides:
            for shape in slide.shapes:
                if hasattr(shape, "text"):
                    text.append(shape.text)
        return "\n".join(text)
    except Exception as e:
        logger.error(f"Error processing PPT file: {str(e)}")
        return f"Error processing PPT file: {str(e)}"


def text_to_speech(text: str, lang: str):
    lang_map = {"English": "en", "Tamil": "ta", "Hindi": "hi"}
    lang_code = lang_map.get(lang, "en")
    tts = gTTS(text=text, lang=lang_code)
    audio_file = "temp_audio.mp3"
    tts.save(audio_file)
    with open(audio_file, "rb") as f:
        audio_bytes = f.read()
    os.remove(audio_file)
    st.session_state.audio_base64 = base64.b64encode(audio_bytes).decode()


def generate_openai_tts(text: str, voice: str):
    try:
        response = client.audio.create(model="tts-1", input=text, voice=voice)
        audio_bytes = response.content
        st.session_state.audio_base64 = base64.b64encode(audio_bytes).decode()
    except openai.APIError as e:
        logger.error(f"OpenAI API error: {e}")
        raise
    except Exception as e:
        logger.error(f"Error generating TTS with OpenAI: {e}")
        raise


def translate_text(text: str, target_lang: str) -> str:
    if target_lang == "English":
        return text
    translator = GoogleTranslator(source="auto", target=target_lang)
    return translator.translate(text)


def update_token_count(tokens: int):
    st.session_state.total_tokens = (
        getattr(st.session_state, "total_tokens", 0) + tokens
    )
    st.session_state.total_cost = (
        getattr(st.session_state, "total_cost", 0) + tokens * 0.0001
    )


def export_chat(format: str) -> str:
    # Check if chat history is available
    if not st.session_state.messages:
        st.warning("No chat history available to export.")
        logger.warning("Export attempted with no chat messages.")
        return None

    # Prepare chat history for export
    chat_history = "\n\n".join(
        [
            f"**{m['role'].capitalize()}:** {m['content']}"
            for m in st.session_state.messages
        ]
    ).strip()  # Remove any leading/trailing whitespace

    # Create export file
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    filename = f"chat_exports/chat_history_{timestamp}.{format}"
    os.makedirs("chat_exports", exist_ok=True)

    try:
        # Write to the appropriate file format
        if format == "md":
            with open(filename, "w", encoding="utf-8") as f:
                f.write(chat_history)
        elif format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, chat_history)
            pdf.output(filename)
        elif format == "txt":
            with open(filename, "w", encoding="utf-8") as f:
                f.write(chat_history)
        elif format == "docx":
            from docx import Document

            doc = Document()
            doc.add_paragraph(chat_history)
            doc.save(filename)
        elif format == "json":
            chat_data = [
                {"role": msg["role"], "content": msg["content"]}
                for msg in st.session_state.messages
            ]
            with open(filename, "w", encoding="utf-8") as f:
                json.dump(chat_data, f, indent=4)

        logger.info(f"Chat exported successfully to {filename}")
        return filename

    except Exception as e:
        logger.error(f"Error exporting chat: {e}", exc_info=True)
        st.error(f"An error occurred while exporting the chat: {e}")
        return None

# main.py
import streamlit as st
import asyncio
import os
from datetime import datetime
from functools import lru_cache
from PIL import Image
from io import StringIO
from fpdf import FPDF

# Import utility functions from utils.py
from utils import (
    perform_ocr,
    process_excel_file,
    process_ppt_file,
    get_max_token_limit,
    get_model_options,
    get_api_client,
    stream_llm_response,
    process_chat_input,
    handle_file_upload,
    export_chat,
)